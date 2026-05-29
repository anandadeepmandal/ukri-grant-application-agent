import datetime
import hashlib
import io
import ipaddress
import re
import socket
from pathlib import Path
from urllib.parse import urljoin, urlparse

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel, Field

APP_NAME = "UKRI Grant Application Agent"
DEVELOPER = "a.mandal@bham.ac.uk"
BASE_DIR = Path(__file__).resolve().parent
HEADERS = {"User-Agent": "UKRI Grant Application Agent/1.0; contact: a.mandal@bham.ac.uk"}

STOPWORDS = set("""
about above across after again against almost along also although always among another around because before being between both could does doing during each either from further have having here into itself just more most other over same should some such than that their there these they this those through under until very were what when where which while with within would research university professor school department publications profile project projects paper papers using based work working publication academic staff student students
""".split())

DOMAIN_PHRASES = [
    "artificial intelligence", "machine learning", "generative ai", "large language model", "data science",
    "financial technology", "fintech", "digital finance", "sustainable finance", "green finance", "climate finance",
    "esg", "responsible ai", "ai governance", "blockchain", "crypto", "digital assets", "risk management",
    "forecasting", "banking", "financial regulation", "consumer protection", "accountability", "digital literacy",
    "skills", "workforce", "healthcare", "nhs", "public health", "one health", "water energy nexus", "sustainability",
    "climate risk", "net zero", "creative economy", "cultural assets", "knowledge exchange", "commercialisation",
    "local authority", "social science", "interdisciplinary", "policy", "innovation", "business engagement", "impact"
]

THEMES = {
    "AI and Data": ["artificial intelligence", "machine learning", "generative ai", "data science", "responsible ai", "ai governance"],
    "FinTech and Finance": ["fintech", "digital finance", "banking", "risk management", "forecasting", "financial regulation"],
    "Sustainability and Climate": ["esg", "sustainable finance", "green finance", "climate finance", "sustainability", "climate risk", "net zero"],
    "Health and Wellbeing": ["healthcare", "nhs", "public health", "one health", "wellbeing"],
    "Policy and Society": ["policy", "social science", "local authority", "accountability", "consumer protection"],
    "Innovation and Impact": ["innovation", "knowledge exchange", "commercialisation", "business engagement", "impact"],
    "Creative and Cultural Economy": ["creative economy", "cultural assets", "digital assets"]
}

class SearchRequest(BaseModel):
    profile_url: str | None = None
    profile_text: str | None = None
    extra_keywords: str | None = None
    max_results: int = Field(default=12, ge=1, le=25)

class TopicRequest(BaseModel):
    profile_url: str | None = None
    profile_text: str | None = None
    extra_keywords: str | None = None
    opportunity_url: str
    opportunity_title: str | None = None
    opportunity_summary: str | None = None

class ApplicationRequest(BaseModel):
    profile_url: str | None = None
    profile_text: str | None = None
    extra_keywords: str | None = None
    opportunity: dict
    selected_topic: dict
    applicant_name: str | None = "Applicant"
    institution: str | None = "University of Birmingham"
    role: str | None = "Principal Investigator"
    partners: str | None = ""
    project_duration_months: str | None = "24"
    budget: str | None = "To be confirmed after costing"
    call_deadline: str | None = ""

app = FastAPI(title=APP_NAME, version="1.0")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_credentials=False, allow_methods=["*"], allow_headers=["*"])

@app.get("/")
def home():
    index_file = BASE_DIR / "index.html"
    if not index_file.exists():
        raise HTTPException(status_code=500, detail=f"Frontend file not found at {index_file}")
    return FileResponse(index_file)

@app.get("/health")
def health():
    return {"status": "ok", "app": APP_NAME, "developer": DEVELOPER}

@app.post("/api/search")
def search(req: SearchRequest):
    profile = build_profile(req.profile_url, req.profile_text, req.extra_keywords)
    keywords = extract_keywords(profile, req.extra_keywords)
    themes = infer_themes(keywords)
    queries = build_queries(keywords, themes)
    opportunities, seen = [], set()
    for query in queries:
        for item in search_ukri(query):
            if item["url"] not in seen:
                seen.add(item["url"])
                opportunities.append(item)
    matches = rank_opportunities(keywords, themes, opportunities)[:req.max_results]
    return {"profile_keywords": keywords, "themes": themes, "queries_used": queries, "matches": matches}

@app.post("/api/topics")
def topics(req: TopicRequest):
    profile = build_profile(req.profile_url, req.profile_text, req.extra_keywords)
    keywords = extract_keywords(profile, req.extra_keywords)
    themes = infer_themes(keywords)
    details = fetch_ukri_details(req.opportunity_url)
    opportunity = {"title": req.opportunity_title or extract_page_title(details), "url": req.opportunity_url, "summary": (details or req.opportunity_summary or "")[:3500]}
    return {"profile_keywords": keywords, "themes": themes, "opportunity": opportunity, "topics": propose_topics(keywords, themes, opportunity)}

@app.post("/api/application")
def application(req: ApplicationRequest):
    profile = build_profile(req.profile_url, req.profile_text, req.extra_keywords)
    keywords = extract_keywords(profile, req.extra_keywords)
    themes = infer_themes(keywords)
    data = build_docx(req, keywords, themes)
    filename = slugify(req.selected_topic.get("title", "ukri_application"))[:80] + "_application_draft.docx"
    return StreamingResponse(io.BytesIO(data), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f'attachment; filename="{filename}"'})

def build_profile(url, text, extra):
    parts = []
    if text and text.strip():
        parts.append(text.strip())
    if url and url.strip():
        parts.append(fetch_public_text(url.strip()))
    if extra and extra.strip():
        parts.append(extra.strip())
    profile = "\n\n".join(parts).strip()
    if not profile:
        raise HTTPException(status_code=400, detail="Please provide a profile URL or pasted profile text.")
    return profile[:80000]

def safe_url(url):
    p = urlparse(url)
    if p.scheme not in {"http", "https"} or not p.hostname:
        return False
    if p.hostname.lower() in {"localhost", "127.0.0.1", "0.0.0.0"}:
        return False
    try:
        for addr in socket.getaddrinfo(p.hostname, None):
            ip = ipaddress.ip_address(addr[4][0])
            if ip.is_private or ip.is_loopback or ip.is_link_local or ip.is_multicast or ip.is_reserved:
                return False
    except Exception:
        return False
    return True

def fetch_public_text(url):
    if not safe_url(url):
        raise HTTPException(status_code=400, detail="The profile URL must be a public http or https URL.")
    try:
        r = requests.get(url, headers=HEADERS, timeout=20)
        r.raise_for_status()
    except requests.RequestException as e:
        raise HTTPException(status_code=400, detail=f"Could not fetch profile URL. Paste profile text instead. Error: {e}")
    return html_to_text(r.text)

def html_to_text(html):
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header", "noscript", "svg"]):
        tag.extract()
    title = soup.title.get_text(" ", strip=True) if soup.title else ""
    body = re.sub(r"\s+", " ", soup.get_text(" ", strip=True))
    return (title + "\n" + body).strip()[:80000]

def extract_keywords(text, extra=None, limit=18):
    t = text.lower()
    scores = {}
    for phrase in DOMAIN_PHRASES:
        c = t.count(phrase)
        if c:
            scores[phrase] = scores.get(phrase, 0) + c * 8
    tokens = [x.replace("-", " ") for x in re.findall(r"[a-zA-Z][a-zA-Z\-]{3,}", t) if x not in STOPWORDS and len(x) > 3]
    for token in tokens:
        scores[token] = scores.get(token, 0) + 1
    for i in range(len(tokens) - 1):
        bg = f"{tokens[i]} {tokens[i+1]}"
        if all(p not in STOPWORDS for p in bg.split()):
            scores[bg] = scores.get(bg, 0) + 2
    if extra:
        for kw in re.split(r"[,;\n]", extra):
            kw = kw.strip().lower()
            if kw:
                scores[kw] = scores.get(kw, 0) + 12
    out = []
    for kw, _ in sorted(scores.items(), key=lambda x: x[1], reverse=True):
        if len(out) >= limit:
            break
        if len(kw) > 2 and not any(kw in e or e in kw for e in out):
            out.append(kw)
    return out

def infer_themes(keywords):
    joined = " ".join(keywords).lower()
    ranked = []
    for theme, terms in THEMES.items():
        score = sum(1 for term in terms if term in joined)
        if score:
            ranked.append((theme, score))
    ranked.sort(key=lambda x: x[1], reverse=True)
    return [x[0] for x in ranked[:4]] or ["Interdisciplinary Research", "Innovation and Impact"]

def build_queries(keywords, themes):
    queries = []
    if keywords:
        queries.append(" ".join(keywords[:4]))
    queries += keywords[:7]
    joined = " ".join(keywords + themes).lower()
    if "ai" in joined or "data" in joined:
        queries.append("artificial intelligence data innovation")
    if "finance" in joined or "banking" in joined:
        queries.append("finance innovation regulation")
    if "sustainability" in joined or "climate" in joined or "esg" in joined:
        queries.append("sustainability climate transition")
    if "health" in joined or "nhs" in joined:
        queries.append("health data innovation")
    if "policy" in joined:
        queries.append("policy evidence innovation")
    queries.append("interdisciplinary research innovation")
    return list(dict.fromkeys([re.sub(r"\s+", " ", q).strip() for q in queries if q.strip()]))[:10]

def search_ukri(query):
    base = "https://www.ukri.org/opportunity/"
    found = []
    for page in range(1, 3):
        url = base if page == 1 else f"{base}page/{page}/"
        params = [("keywords", query), ("filter_status[]", "open"), ("filter_status[]", "upcoming"), ("filter_submitted", "true"), ("filter_order", "closing_date")]
        try:
            r = requests.get(url, params=params, headers=HEADERS, timeout=20)
            r.raise_for_status()
        except requests.RequestException:
            continue
        items = parse_ukri(r.text)
        if not items:
            break
        found += items
    return found

def parse_ukri(html):
    soup, items, seen = BeautifulSoup(html, "html.parser"), [], set()
    for a in soup.find_all("a", href=True):
        href = urljoin("https://www.ukri.org", a["href"]).split("#")[0].split("?")[0].rstrip("/")
        title = a.get_text(" ", strip=True)
        if not title or len(title) < 8 or "/opportunity/" not in href or href == "https://www.ukri.org/opportunity" or href in seen:
            continue
        seen.add(href)
        box = a
        for _ in range(6):
            if box.parent:
                box = box.parent
            if len(box.get_text(" ", strip=True)) > 260:
                break
        text = re.sub(r"\s+", " ", box.get_text(" ", strip=True))
        items.append({"id": hashlib.md5(href.encode()).hexdigest()[:12], "title": title, "url": href + "/", "summary": text[:1200], "status": field(text, r"Opportunity status:\s*(Open|Upcoming|Closed)"), "funders": field(text, r"Funders?:\s*(.*?)(?:Funding type:|Co-funders:|Total fund:|Maximum award:|Publication date:|Opening date:|Closing date:|$)"), "funding_type": field(text, r"Funding type:\s*(.*?)(?:Total fund:|Maximum award:|Award:|Publication date:|Opening date:|Closing date:|$)"), "award": field(text, r"(?:Total fund|Maximum award|Award):\s*(.*?)(?:Publication date:|Opening date:|Closing date:|$)"), "closing_date": field(text, r"Closing date:\s*(.*?)(?:$|Publication date:|Opening date:)")})
    return items

def field(text, pattern):
    m = re.search(pattern, text, flags=re.I)
    return re.sub(r"\s+", " ", m.group(1)).strip()[:240] if m else ""

def rank_opportunities(keywords, themes, items):
    ranked = []
    for item in items:
        text = f"{item.get('title','')} {item.get('summary','')} {item.get('funders','')} {item.get('funding_type','')}".lower()
        matched = [kw for kw in keywords if kw.lower() in text]
        theme_hits = [th for th in themes if any(w.lower() in text for w in th.split())]
        score = min(98, 25 + min(len(matched) * 8, 45) + min(len(theme_hits) * 8, 20) + (6 if item.get("status", "").lower() == "open" else 0) + (4 if item.get("closing_date") else 0))
        x = dict(item)
        x["match_score"] = score
        x["matched_terms"] = matched[:10]
        x["theme_hits"] = theme_hits[:5]
        x["why_match"] = ("It overlaps with your profile terms: " + ", ".join(matched[:6]) + ".") if matched else "This opportunity may be relevant based on broader thematic similarity."
        ranked.append(x)
    return sorted(ranked, key=lambda x: x["match_score"], reverse=True)

def fetch_ukri_details(url):
    p = urlparse(url)
    if p.scheme not in {"http", "https"} or not p.hostname or not p.hostname.lower().endswith("ukri.org"):
        raise HTTPException(status_code=400, detail="The selected opportunity must be from ukri.org.")
    try:
        r = requests.get(url, headers=HEADERS, timeout=20)
        r.raise_for_status()
    except requests.RequestException as e:
        raise HTTPException(status_code=400, detail=f"Could not fetch the UKRI page. Error: {e}")
    return html_to_text(r.text)

def extract_page_title(text):
    parts = [x.strip() for x in text.splitlines() if x.strip()]
    return parts[0][:180] if parts else "Selected UKRI opportunity"

def context_for(keywords, themes):
    joined = " ".join(keywords + themes).lower()
    if "finance" in joined or "fintech" in joined or "banking" in joined:
        return "responsible digital finance and financial innovation"
    if "health" in joined or "nhs" in joined:
        return "health data innovation and public value"
    if "climate" in joined or "sustainability" in joined or "esg" in joined:
        return "sustainable transition, ESG and climate informed decision making"
    if "creative" in joined or "cultural" in joined:
        return "creative and cultural data assets"
    if "local authority" in joined or "policy" in joined:
        return "policy innovation and public sector decision making"
    if "ai" in joined or "machine learning" in joined or "data science" in joined:
        return "responsible AI, data science and digital transformation"
    return "interdisciplinary innovation and research impact"

def methods_for(keywords):
    joined = " ".join(keywords).lower()
    methods = []
    if "machine learning" in joined or "ai" in joined or "data" in joined:
        methods += ["machine learning", "natural language processing", "explainable AI"]
    if "finance" in joined or "risk" in joined or "forecasting" in joined:
        methods += ["econometric modelling", "risk analytics", "forecasting"]
    if "policy" in joined or "social" in joined or "skills" in joined:
        methods += ["mixed methods", "stakeholder interviews", "policy analysis"]
    if "health" in joined:
        methods += ["health data analytics", "case study analysis", "implementation evaluation"]
    if "sustainability" in joined or "climate" in joined:
        methods += ["impact measurement", "scenario analysis", "sustainability assessment"]
    return list(dict.fromkeys(methods or ["mixed methods", "stakeholder analysis", "impact evaluation"]))[:4]

def partners_for(themes, keywords):
    joined = " ".join(themes + keywords).lower()
    partners = ["academic co investigators", "industry or policy stakeholders", "knowledge exchange partners"]
    if "finance" in joined or "banking" in joined:
        partners += ["banks or financial institutions", "financial regulators or professional bodies", "FinTech firms"]
    if "health" in joined or "nhs" in joined:
        partners += ["NHS trusts", "public health bodies", "health innovation networks"]
    if "local authority" in joined or "policy" in joined:
        partners += ["local authorities", "policy units", "community organisations"]
    if "creative" in joined or "cultural" in joined:
        partners += ["museums and cultural organisations", "creative businesses", "rights and licensing specialists"]
    if "sustainability" in joined or "climate" in joined or "esg" in joined:
        partners += ["sustainability networks", "environmental organisations", "corporate ESG teams"]
    return list(dict.fromkeys(partners))[:8]

def propose_topics(keywords, themes, opportunity):
    ctx, methods, partners = context_for(keywords, themes), methods_for(keywords), partners_for(themes, keywords)
    call = opportunity.get("title", "selected UKRI opportunity")
    templates = [("AI enabled evidence and decision support for {ctx}", "Builds a decision support framework that turns complex data into actionable insight for researchers, policy makers, organisations and public stakeholders."), ("Responsible innovation and governance framework for {ctx}", "Develops a responsible innovation model addressing accountability, fairness, transparency, adoption risks and implementation pathways."), ("Data driven measurement and impact evaluation of {ctx}", "Creates an empirical framework for measuring outcomes, evaluating interventions and producing evidence for funders, regulators and practitioners."), ("Cross sector knowledge exchange platform for {ctx}", "Focuses on stakeholder engagement, translational outputs, demonstrators, workshops, training materials and policy facing impact."), ("Digital capability, skills and adoption pathway for {ctx}", "Links research expertise to training, digital adoption, organisational change, workforce capability and long term sustainability."), ("Risk, resilience and sustainable transformation in {ctx}", "Frames the project around risk identification, resilience building, sustainability outcomes and scalable transformation.")]
    out = []
    for i, (title, rationale) in enumerate(templates, 1):
        title = title.format(ctx=ctx)
        out.append({"id": f"topic-{i}", "title": title, "fit_score": max(95 - i * 4, 72), "call_alignment": f"Designed for the call: {call}. The topic uses the applicant's strengths in {', '.join(keywords[:5])}.", "rationale": rationale, "research_questions": [f"What are the main barriers and enablers shaping {ctx}?", f"How can {', '.join(methods[:2])} be used to produce reliable and decision relevant evidence?", "What governance, adoption and impact mechanisms are needed for responsible implementation?"], "methodology": ["Rapid evidence review and stakeholder mapping", "Data collection through documents, interviews, workshops or open datasets", f"Analytical work using {', '.join(methods)}", "Validation through expert review, case study testing and policy or practitioner feedback"], "work_packages": ["WP1, Scoping, literature review and stakeholder mapping", "WP2, Data collection, model design and analytical framework", "WP3, Validation, demonstrator development and knowledge exchange", "WP4, Impact, dissemination, policy translation and future funding pathway"], "impact_pathway": "The project can produce academic outputs, stakeholder workshops, a policy or practice brief, reusable datasets or tools where appropriate, and a larger follow on funding application.", "suggested_partners": partners, "matched_keywords": keywords[:5]})
    return out

def build_docx(req, keywords, themes):
    doc = Document()
    style_doc(doc)
    opp, topic = req.opportunity, req.selected_topic
    title = topic.get("title", "UKRI Application Draft")
    call = opp.get("title", "Selected UKRI opportunity")
    add_title(doc, "UKRI Funding Application Draft")
    add_subtitle(doc, f"Generated by {APP_NAME}, developed by {DEVELOPER}")
    meta(doc, [("Generated on", datetime.date.today().strftime("%d %B %Y")), ("Applicant", req.applicant_name or "Applicant"), ("Institution", req.institution or "University of Birmingham"), ("Applicant role", req.role or "Principal Investigator"), ("Selected UKRI call", call), ("Call URL", opp.get("url", "")), ("Selected project topic", title), ("Duration", f"{req.project_duration_months or '24'} months"), ("Indicative budget", req.budget or "To be confirmed after costing"), ("Call deadline", req.call_deadline or opp.get("closing_date", "To be confirmed"))])
    section(doc, "1. Project title"); doc.add_paragraph(title)
    section(doc, "2. Plain English summary"); doc.add_paragraph(f"This project responds to {call} by developing a focused programme of work on {title.lower()}. The proposal draws on the applicant profile themes of {', '.join(themes)} and expertise in {', '.join(keywords[:8])}. It will generate rigorous evidence, practical outputs and a clear pathway to academic, policy and stakeholder impact.")
    section(doc, "3. Fit with the selected UKRI opportunity"); doc.add_paragraph(topic.get("call_alignment", "The project has been designed to align with the selected UKRI funding opportunity.")); doc.add_paragraph("The official UKRI call page must be checked for eligibility, scope, funding limits, required attachments, cost rules and submission route.")
    section(doc, "4. Applicant expertise and track record"); doc.add_paragraph(f"The applicant profile indicates expertise in {', '.join(keywords[:10])}. This provides a credible foundation for the proposed project because it combines subject knowledge, methodological capability and external engagement potential.")
    section(doc, "5. Research problem and rationale"); doc.add_paragraph(topic.get("rationale", "The project addresses a timely research, innovation and impact challenge."))
    section(doc, "6. Aim and objectives"); doc.add_paragraph(f"Overall aim: To develop a rigorous and impact oriented research programme on {title.lower()} aligned with the priorities of the selected UKRI opportunity."); bullets(doc, ["Map the research, policy and practice landscape.", "Develop a robust analytical or conceptual framework.", "Collect and analyse evidence using appropriate methods.", "Co design outputs with stakeholders and partners.", "Produce academic, policy and practice facing outputs."])
    section(doc, "7. Research questions"); bullets(doc, topic.get("research_questions", []))
    section(doc, "8. Methodology"); bullets(doc, topic.get("methodology", []))
    section(doc, "9. Work packages"); bullets(doc, topic.get("work_packages", []))
    section(doc, "10. Partnership and stakeholder engagement")
    if req.partners: doc.add_paragraph(f"Proposed named partners: {req.partners}")
    bullets(doc, topic.get("suggested_partners", []))
    section(doc, "11. Outputs and deliverables"); bullets(doc, ["Academic paper or working paper", "Policy or practitioner briefing", "Stakeholder workshop or knowledge exchange event", "Technical or methodological appendix, where appropriate", "Roadmap for follow on funding"])
    section(doc, "12. Impact pathway"); doc.add_paragraph(topic.get("impact_pathway", "The project will produce academic, policy and practice impact."))
    section(doc, "13. Equality, diversity and inclusion"); doc.add_paragraph("The project will apply inclusive recruitment, engagement and dissemination practices.")
    section(doc, "14. Ethics, data management and responsible research"); doc.add_paragraph("The project will follow institutional ethics, data protection and responsible research requirements.")
    section(doc, "15. Risk management"); risk_table(doc)
    section(doc, "16. Indicative budget narrative"); doc.add_paragraph(f"The indicative budget is {req.budget or 'to be confirmed after costing'}. Eligible cost categories should be checked against the official UKRI call guidance.")
    section(doc, "17. Indicative timeline"); timeline_table(doc, req.project_duration_months or "24")
    section(doc, "18. Submission checklist"); bullets(doc, ["Confirm applicant and institutional eligibility", "Check the official UKRI call page for deadline and submission route", "Confirm funding limits and costing rules", "Prepare case for support and required attachments", "Secure partner letters where required", "Complete institutional approval workflow"])
    section(doc, "19. Matched profile keywords"); bullets(doc, keywords[:15])
    section(doc, "20. Final note"); doc.add_paragraph("This is a structured application draft. It must be reviewed, costed and adapted to the official UKRI application form before submission.")
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()

def style_doc(doc):
    s = doc.sections[0]
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(0.7)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

def add_title(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(18)

def add_subtitle(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.italic = True; r.font.size = Pt(10)

def section(doc, text):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = True; r.font.size = Pt(13)

def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet"); p.add_run(str(item))

def meta(doc, rows):
    table = doc.add_table(rows=0, cols=2); table.style = "Table Grid"
    for k, v in rows:
        cells = table.add_row().cells; cells[0].text = k; cells[1].text = v or ""
        for r in cells[0].paragraphs[0].runs: r.bold = True
    doc.add_paragraph("")

def risk_table(doc):
    risks = [("Call scope mismatch", "Moderate", "Review guidance and align objectives."), ("Partner engagement delay", "Moderate", "Identify primary and reserve partners early."), ("Data access constraints", "Moderate", "Use alternative public datasets or case studies."), ("Ethics delay", "Low to moderate", "Begin ethics planning during setup."), ("Impact pathway too broad", "Moderate", "Focus outputs on defined audiences.")]
    table = doc.add_table(rows=1, cols=3); table.style = "Table Grid"
    table.rows[0].cells[0].text = "Risk"; table.rows[0].cells[1].text = "Level"; table.rows[0].cells[2].text = "Mitigation"
    for a, b, c in risks:
        row = table.add_row().cells; row[0].text = a; row[1].text = b; row[2].text = c

def timeline_table(doc, duration):
    try: months = int(re.sub(r"[^0-9]", "", duration) or "24")
    except Exception: months = 24
    periods = [("Months 1 to 6", "Scoping, review, ethics, partner engagement and design"), ("Months 7 to 12", "Data collection, workshops and initial analysis"), ("Months 13 to 18", "Framework, model or case study development and validation"), ("Months 19 to 24", "Impact, dissemination, publications and follow on funding")]
    if months <= 12:
        periods = [("Months 1 to 3", "Scoping, ethics and review"), ("Months 4 to 7", "Data collection and analysis"), ("Months 8 to 10", "Validation and outputs"), ("Months 11 to 12", "Dissemination and follow on plan")]
    table = doc.add_table(rows=1, cols=2); table.style = "Table Grid"
    table.rows[0].cells[0].text = "Period"; table.rows[0].cells[1].text = "Main activities"
    for a, b in periods:
        row = table.add_row().cells; row[0].text = a; row[1].text = b

def slugify(text):
    return re.sub(r"[^A-Za-z0-9]+", "_", text).strip("_").lower() or "ukri_application"
